from docx import Document
from reportlab.pdfgen import canvas
import os

OUTPUT_FOLDER = "generated_resumes"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

candidates = [
    {
        "name": "Arjun Menon",
        "email": "arjun.menon@gmail.com",
        "phone": "+91 9876543211",
        "linkedin": "https://linkedin.com/in/arjunmenon",
        "role": "Junior Software Engineer",
        "location": "Kochi",
        "experience": "1 Year",
        "skills": "Python, SQL, Flask, Git",
        "education": "B.Tech Computer Science",
        "type": "docx"
    },
    {
        "name": "Priya Nair",
        "email": "priya.nair@gmail.com",
        "phone": "+91 9876543212",
        "linkedin": "https://linkedin.com/in/priyanair",
        "role": "UI/UX Designer",
        "location": "Bangalore",
        "experience": "3 Years",
        "skills": "Figma, Adobe XD, User Research, Wireframing",
        "education": "B.Des Design",
        "type": "pdf"
    },
    {
        "name": "Rahul Sharma",
        "email": "rahul.sharma@gmail.com",
        "phone": "+91 9876543213",
        "linkedin": "https://linkedin.com/in/rahulsharma",
        "role": "Project Manager",
        "location": "Pune",
        "experience": "8 Years",
        "skills": "Agile, Scrum, Jira, Risk Management",
        "education": "MBA Project Management",
        "type": "pdf"
    },
    {
        "name": "Sneha Reddy",
        "email": "sneha.reddy@gmail.com",
        "phone": "+91 9876543214",
        "linkedin": "https://linkedin.com/in/snehareddy",
        "role": "Product Manager",
        "location": "Hyderabad",
        "experience": "6 Years",
        "skills": "Product Strategy, Analytics, Roadmaps, SQL",
        "education": "MBA",
        "type": "docx"
    },
    {
        "name": "Vikram Patel",
        "email": "vikram.patel@gmail.com",
        "phone": "+91 9876543215",
        "linkedin": "https://linkedin.com/in/vikrampatel",
        "role": "Backend Developer",
        "location": "Mumbai",
        "experience": "5 Years",
        "skills": "Java, Spring Boot, PostgreSQL, Docker",
        "education": "B.Tech IT",
        "type": "docx"
    },
    {
        "name": "Akash Nambiar",
        "email": "akash.nambiar@gmail.com",
        "phone": "+91 9876543226",
        "linkedin": "https://linkedin.com/in/akashnambiar",
        "role": "DevOps Engineer",
        "location": "Kochi",
        "experience": "5 Years",
        "skills": "Docker, Kubernetes, AWS, Jenkins",
        "education": "B.Tech Computer Science",
        "type": "pdf"
    },
    {
        "name": "Pooja Iyer",
        "email": "pooja.iyer@gmail.com",
        "phone": "+91 9876543227",
        "linkedin": "https://linkedin.com/in/poojaiyer",
        "role": "HR Executive",
        "location": "Bangalore",
        "experience": "2 Years",
        "skills": "Recruitment, HRMS, Onboarding, Communication",
        "education": "MBA HR",
        "type": "docx"
    },
    {
        "name": "Nikhil Bhat",
        "email": "nikhil.bhat@gmail.com",
        "phone": "+91 9876543228",
        "linkedin": "https://linkedin.com/in/nikhilbhat",
        "role": "Data Analyst",
        "location": "Pune",
        "experience": "3 Years",
        "skills": "Python, SQL, Power BI, Excel",
        "education": "B.Tech IT",
        "type": "pdf"
    },
    {
        "name": "Farah Ali",
        "email": "farah.ali@gmail.com",
        "phone": "+91 9876543229",
        "linkedin": "https://linkedin.com/in/farahali",
        "role": "Product Manager",
        "location": "Hyderabad",
        "experience": "6 Years",
        "skills": "Product Strategy, Agile, Jira, Roadmapping",
        "education": "MBA",
        "type": "docx"
    },
    {
        "name": "Kiran Thomas",
        "email": "kiran.thomas@gmail.com",
        "phone": "+91 9876543230",
        "linkedin": "https://linkedin.com/in/kiranthomas",
        "role": "QA Engineer",
        "location": "Chennai",
        "experience": "4 Years",
        "skills": "Selenium, Postman, API Testing, Automation",
        "education": "B.Tech Computer Science",
        "type": "pdf"
    },
    {
        "name": "Deepa Nair",
        "email": "deepa.nair@gmail.com",
        "phone": "+91 9876543231",
        "linkedin": "https://linkedin.com/in/deepanair",
        "role": "Operations Executive",
        "location": "Kochi",
        "experience": "1 Year",
        "skills": "Operations, Reporting, Excel, Documentation",
        "education": "B.Com",
        "type": "docx"
    },
    {
        "name": "Ritesh Gupta",
        "email": "ritesh.gupta@gmail.com",
        "phone": "+91 9876543232",
        "linkedin": "https://linkedin.com/in/riteshgupta",
        "role": "Senior Software Engineer",
        "location": "Bangalore",
        "experience": "8 Years",
        "skills": "Java, Spring Boot, Microservices, AWS",
        "education": "B.Tech Computer Science",
        "type": "pdf"
    },
    {
        "name": "Ananya Roy",
        "email": "ananya.roy@gmail.com",
        "phone": "+91 9876543233",
        "linkedin": "https://linkedin.com/in/ananyaroy",
        "role": "Digital Marketing Specialist",
        "location": "Mumbai",
        "experience": "3 Years",
        "skills": "SEO, Google Analytics, Content Marketing",
        "education": "MBA Marketing",
        "type": "docx"
    },
    {
        "name": "Mohammed Irfan",
        "email": "mohammed.irfan@gmail.com",
        "phone": "+91 9876543234",
        "linkedin": "https://linkedin.com/in/mohammedirfan",
        "role": "Network Engineer",
        "location": "Hyderabad",
        "experience": "5 Years",
        "skills": "Cisco, Routing, Switching, Network Security",
        "education": "B.Tech ECE",
        "type": "pdf"
    },
    {
        "name": "Shruti Pillai",
        "email": "shruti.pillai@gmail.com",
        "phone": "+91 9876543235",
        "linkedin": "https://linkedin.com/in/shrutipillai",
        "role": "UI Designer",
        "location": "Trivandrum",
        "experience": "2 Years",
        "skills": "Figma, Adobe XD, Prototyping, User Research",
        "education": "B.Des Design",
        "type": "docx"
    },
    {
        "name": "Varun Sreedhar",
        "email": "varun.sreedhar@gmail.com",
        "phone": "+91 9876543236",
        "linkedin": "https://linkedin.com/in/varunsreedhar",
        "role": "Business Analyst",
        "location": "Bangalore",
        "experience": "4 Years",
        "skills": "SQL, Power BI, Excel, Requirement Gathering",
        "education": "MBA",
        "type": "pdf"
    },
    {
        "name": "Gayathri Menon",
        "email": "gayathri.menon@gmail.com",
        "phone": "+91 9876543237",
        "linkedin": "https://linkedin.com/in/gayathrimenon",
        "role": "Customer Success Manager",
        "location": "Chennai",
        "experience": "5 Years",
        "skills": "CRM, Client Management, Communication",
        "education": "MBA",
        "type": "docx"
    },
    {
        "name": "Naveen Raj",
        "email": "naveen.raj@gmail.com",
        "phone": "+91 9876543238",
        "linkedin": "https://linkedin.com/in/naveenraj",
        "role": "Cybersecurity Analyst",
        "location": "Pune",
        "experience": "3 Years",
        "skills": "SOC Monitoring, SIEM, Incident Response",
        "education": "B.Tech Computer Science",
        "type": "pdf"
    },
    {
        "name": "Reshma George",
        "email": "reshma.george@gmail.com",
        "phone": "+91 9876543239",
        "linkedin": "https://linkedin.com/in/reshmageorge",
        "role": "Project Coordinator",
        "location": "Kochi",
        "experience": "2 Years",
        "skills": "Agile, Documentation, Project Tracking",
        "education": "MBA",
        "type": "docx"
    },
    {
        "name": "Suresh Kumar",
        "email": "suresh.kumar@gmail.com",
        "phone": "+91 9876543240",
        "linkedin": "https://linkedin.com/in/sureshkumar",
        "role": "Sales Manager",
        "location": "Delhi",
        "experience": "7 Years",
        "skills": "CRM, Lead Generation, Negotiation, Sales Strategy",
        "education": "BBA",
        "type": "pdf"
    },
    {
        "name": "Meera Krishnan",
        "email": "meera.krishnan@gmail.com",
        "phone": "+91 9876543216",
        "linkedin": "https://linkedin.com/in/meerakrishnan",
        "role": "Graphic Designer",
        "location": "Trivandrum",
        "experience": "4 Years",
        "skills": "Photoshop, Illustrator, Branding, Canva",
        "education": "B.Des",
        "type": "pdf"
    },
    {
        "name": "Ankit Verma",
        "email": "ankit.verma@gmail.com",
        "phone": "+91 9876543217",
        "linkedin": "https://linkedin.com/in/ankitverma",
        "role": "Sales Executive",
        "location": "Delhi",
        "experience": "2 Years",
        "skills": "CRM, Lead Generation, Negotiation",
        "education": "BBA",
        "type": "docx"
    },
    {
        "name": "Neha Joseph",
        "email": "neha.joseph@gmail.com",
        "phone": "+91 9876543218",
        "linkedin": "https://linkedin.com/in/nehajoseph",
        "role": "Operations Manager",
        "location": "Chennai",
        "experience": "7 Years",
        "skills": "Supply Chain, Process Optimization, ERP",
        "education": "MBA Operations",
        "type": "docx"
    },
    {
        "name": "Aditya Rao",
        "email": "aditya.rao@gmail.com",
        "phone": "+91 9876543219",
        "linkedin": "https://linkedin.com/in/adityarao",
        "role": "Full Stack Developer",
        "location": "Bangalore",
        "experience": "4 Years",
        "skills": "React, Node.js, MongoDB, AWS",
        "education": "B.Tech CSE",
        "type": "docx"
    },
    {
        "name": "Kavya Das",
        "email": "kavya.das@gmail.com",
        "phone": "+91 9876543220",
        "linkedin": "https://linkedin.com/in/kavyadas",
        "role": "QA Automation Engineer",
        "location": "Hyderabad",
        "experience": "3 Years",
        "skills": "Selenium, Python, Test Automation",
        "education": "B.Tech",
        "type": "pdf"
    },
    {
        "name": "Rohan Iyer",
        "email": "rohan.iyer@gmail.com",
        "phone": "+91 9876543221",
        "linkedin": "https://linkedin.com/in/rohaniyer",
        "role": "DevOps Engineer",
        "location": "Bangalore",
        "experience": "6 Years",
        "skills": "AWS, Docker, Kubernetes, Linux",
        "education": "B.Tech",
        "type": "docx"
    },
    {
        "name": "Aisha Khan",
        "email": "aisha.khan@gmail.com",
        "phone": "+91 9876543222",
        "linkedin": "https://linkedin.com/in/aishakhan",
        "role": "UI Designer",
        "location": "Kochi",
        "experience": "1 Year",
        "skills": "Figma, Wireframing, Prototyping",
        "education": "B.Des",
        "type": "docx"
    },
    {
        "name": "Sanjay Nair",
        "email": "sanjay.nair@gmail.com",
        "phone": "+91 9876543223",
        "linkedin": "https://linkedin.com/in/sanjaynair",
        "role": "Data Engineer",
        "location": "Pune",
        "experience": "5 Years",
        "skills": "Python, SQL, Spark, AWS",
        "education": "B.Tech",
        "type": "pdf"
    },
    {
        "name": "Divya Menon",
        "email": "divya.menon@gmail.com",
        "phone": "+91 9876543224",
        "linkedin": "https://linkedin.com/in/divyamenon",
        "role": "Business Analyst",
        "location": "Chennai",
        "experience": "4 Years",
        "skills": "SQL, Power BI, Excel, Requirements Gathering",
        "education": "MBA",
        "type": "docx"
    },
    {
        "name": "Harish Kumar",
        "email": "harish.kumar@gmail.com",
        "phone": "+91 9876543225",
        "linkedin": "https://linkedin.com/in/harishkumar",
        "role": "Senior Software Engineer",
        "location": "Mumbai",
        "experience": "10 Years",
        "skills": "Python, Django, PostgreSQL, AWS",
        "education": "B.Tech",
        "type": "pdf"
    }
]

def create_docx(candidate):
    doc = Document()

    doc.add_heading(candidate["name"], level=1)
    doc.add_paragraph(f"Email: {candidate['email']}")
    doc.add_paragraph(f"Phone: {candidate['phone']}")
    doc.add_paragraph(f"LinkedIn: {candidate['linkedin']}")
    doc.add_paragraph(f"Role: {candidate['role']}")
    doc.add_paragraph(f"Location: {candidate['location']}")
    doc.add_paragraph(f"Experience: {candidate['experience']}")
    doc.add_paragraph(f"Skills: {candidate['skills']}")
    doc.add_paragraph(f"Education: {candidate['education']}")

    filename = os.path.join(
        OUTPUT_FOLDER,
        candidate["name"].replace(" ", "_") + ".docx"
    )

    doc.save(filename)

def create_pdf(candidate):
    filename = os.path.join(
        OUTPUT_FOLDER,
        candidate["name"].replace(" ", "_") + ".pdf"
    )

    c = canvas.Canvas(filename)

    y = 800

    lines = [
        candidate["name"],
        f"Email: {candidate['email']}",
        f"Phone: {candidate['phone']}",
        f"LinkedIn: {candidate['linkedin']}",
        f"Role: {candidate['role']}",
        f"Location: {candidate['location']}",
        f"Experience: {candidate['experience']}",
        f"Skills: {candidate['skills']}",
        f"Education: {candidate['education']}"
    ]

    for line in lines:
        c.drawString(50, y, line)
        y -= 20

    c.save()

for candidate in candidates:
    if candidate["type"] == "docx":
        create_docx(candidate)
    else:
        create_pdf(candidate)

print("Resumes generated successfully!")