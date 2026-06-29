from flask import Flask, render_template, request ,send_from_directory, session
from werkzeug.utils import secure_filename
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from functools import wraps
import sqlite3
import os
import re
from pypdf import PdfReader
from docx import Document
from dotenv import load_dotenv
from groq import Groq
from database import save_candidate, save_retry_job, get_pending_retry_jobs, mark_retry_job_completed
from flask import redirect
import logging
from logging.handlers import RotatingFileHandler
import time

app = Flask(__name__)

handler = RotatingFileHandler(
    "app.log",
    maxBytes=1024 * 1024,
    backupCount=5
)

handler.setLevel(logging.INFO)

app.logger.addHandler(handler)

app.logger.setLevel(logging.INFO)

app.secret_key = "change_this_for_production"

limiter = Limiter(
    key_func=get_remote_address,
    app=app,
    default_limits=[]
)

def login_required(f):

    @wraps(f)
    def decorated_function(*args, **kwargs):

        if not session.get("logged_in"):
            return redirect("/login")

        return f(*args, **kwargs)

    return decorated_function

load_dotenv()

client = Groq(
    api_key=os.getenv("GROQ_API_KEY")
)

UPLOAD_FOLDER = "uploads"
import os
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 5 * 1024 * 1024

@app.route("/login", methods=["GET", "POST"])
def login():

    if request.method == "POST":

        username = request.form.get("username")
        password = request.form.get("password")

        if (
            username == "recruiter"
            and
            password == "password123"
        ):

            session["logged_in"] = True

            return redirect("/")

        return render_template(
            "login.html",
            error="Invalid username or password."
        )

    return render_template("login.html")

@app.route("/logout")
def logout():

    session.clear()

    return redirect("/login")

@app.route("/")
@login_required
def home():

    message = request.args.get("message", "")
    error = request.args.get("error", "")

    conn = sqlite3.connect("candidates.db")
    cursor = conn.cursor()

    cursor.execute("""
    CREATE TABLE IF NOT EXISTS candidates (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        phone TEXT,
        linkedin TEXT,
        filename TEXT,
        skills TEXT,
        experience TEXT,
        job_title TEXT,
        location TEXT
    )
    """)

    cursor.execute("""
    SELECT *
    FROM candidates
    ORDER BY id DESC
    """)

    candidates = cursor.fetchall()

    conn.commit()
    conn.close()

    return render_template(
        "index.html",
        candidates=candidates,
        message=message,
        error=error
    )

def parse_ai_response(response):

    skills = "Not Available"
    experience = "0"
    job_title = "Not Available"
    location = "Not Available"

    if response:

        ai_output = response.choices[0].message.content

        for line in ai_output.split("\n"):

            if line.startswith("SKILLS:"):
                skills = line.replace(
                    "SKILLS:",
                    ""
                ).strip()

            elif line.startswith("EXPERIENCE:"):
                experience = line.replace(
                    "EXPERIENCE:",
                    ""
                ).strip()

            elif line.startswith("JOB_TITLE:"):
                job_title = line.replace(
                    "JOB_TITLE:",
                    ""
                ).strip()

            elif line.startswith("LOCATION:"):
                location = line.replace(
                    "LOCATION:",
                    ""
                ).strip()

    return (
        skills,
        experience,
        job_title,
        location
    )

def analyze_resume(prompt, filename):

    start_time = time.time()

    for attempt in range(3):

        try:

            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            )

            end_time = time.time()

            app.logger.info(
                f"AI API successful: {filename} | Time taken: {end_time - start_time:.2f} seconds"
            )

            return response

        except Exception as e:

            end_time = time.time()

            app.logger.error(
                f"AI API failed: {filename} | Time taken: {end_time - start_time:.2f} seconds | Error: {str(e)}"
            )

            save_retry_job(filename)

            print("GROQ ERROR:", e)

            return None

def process_resume(file_path, filename):

    print(f"Processing resume: {filename}")

    text = ""

    if filename.endswith(".pdf"):

        reader = PdfReader(file_path)

        for page in reader.pages:

            page_text = page.extract_text()

            text += page_text or ""

    elif filename.endswith(".docx"):

        doc = Document(file_path)

        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

    print("Resume text extracted.")

    lines = text.split("\n")

    name = ""

    for line in lines:

        line = line.strip()

        if not line:
            continue

        if "@" in line:
            continue

        if "linkedin" in line.lower():
            continue

        if "github" in line.lower():
            continue

        if len(line) > 40:
            continue

        name = line
        break

    print(f"Candidate extracted from file : {filename}")

    emails = re.findall(
        r'[\w\.-]+@[\w\.-]+\.\w+',
        text
    )

    phones = re.findall(
        r'(\+?\d[\d\s\-]{8,15}\d)',
        text
    )

    linkedin_matches = re.findall(
        r'(?:https?://)?(?:www\.)?linkedin\.com/[^\s]+',
        text,
        re.IGNORECASE
    )

    linkedin = ""

    if linkedin_matches:
        linkedin = linkedin_matches[0]
    else:
        linkedin_short = re.findall(
            r'linkedin/[A-Za-z0-9_-]+',
            text,
            re.IGNORECASE
        )

        if linkedin_short:
            linkedin = linkedin_short[0]

    github = ""

    github_matches = re.findall(
        r'github/[A-Za-z0-9_-]+',
        text,
        re.IGNORECASE
    )

    if github_matches:
        github = github_matches[0]

    email = emails[0] if emails else ""
    phone = phones[0] if phones else ""

    print("Contact details extracted successfully")

    return (
        text,
        name,
        email,
        phone,
        linkedin,
        github,
        emails,
        phones
    )


@app.route("/upload", methods=["POST"])
@login_required
@limiter.limit("30 per minute")
def upload():

    files = request.files.getlist("resumes")
    success_count = 0
    failed_files = []
    failed_reasons = []

    if not files or files[0].filename == "":
        return "No file selected!"

    for file in files:
        try:
            print("Processing File:", file.filename)
            
            filename = secure_filename(file.filename)

            app.logger.info(f"Upload started: {filename}")

            file_path = os.path.join(
                app.config["UPLOAD_FOLDER"],
                filename
            )

            if not (
                file.filename.lower().endswith(".pdf")
                or
                file.filename.lower().endswith(".docx")
            ):
                failed_files.append(file.filename)
                failed_reasons.append(
                    "Only PDF and DOCX files are allowed."
                )
                continue

            file.save(file_path)

            text, name, email, phone, linkedin, github,emails, phones = process_resume(
                file_path,
                filename
            )

            scrubbed_text = text

            if linkedin:
                scrubbed_text = scrubbed_text.replace(
                    linkedin,
                    "[LINKEDIN]"
                )
            
            if github:
                scrubbed_text = scrubbed_text.replace(
                    github,
                    "[GITHUB]"
                )
            for email in emails:
                scrubbed_text = scrubbed_text.replace(
                    email,
                    "[EMAIL]"
                )

            for phone in phones:
                scrubbed_text = scrubbed_text.replace(
                    phone,
                    "[PHONE]"
                )

            scrubbed_text = scrubbed_text.replace(
                "LinkedIn",
                "[LINKEDIN]"
            )
            scrubbed_text = scrubbed_text.replace(
                "GitHub",
                "[GITHUB]"
            )

            print("Resume text scrubbed successfully.")


            prompt = f"""
        Analyze this resume.

        Return ONLY in this exact format:

        SKILLS:<comma separated skills>
        EXPERIENCE:<years>
        JOB_TITLE:<most recent job title>
        LOCATION:<location>

        Resume:

        {scrubbed_text}
        """
            response = analyze_resume(
                prompt,
                filename
            )

            skills, experience, job_title, location = parse_ai_response(
                response
            )

            if response:
                print(f"AI analysis completed: {filename}")
                
                print("AI fields extracted successfully.")
                print("RESPONSE IS NONE =", response is None)

                save_candidate(
                        name,
                        email,
                        phone,
                        linkedin,
                        filename,
                        skills,
                        experience,
                        job_title,
                        location
                )

                success_count += 1
                app.logger.info(f"Upload successful: {filename}")
            else:
                continue
        except Exception as e:
            print("ERROR:", e)

            error_message = str(e)

            app.logger.error(
                f"Upload failed: {file.filename} | Reason: {error_message}"
            )

            if "Stream has ended unexpectedly" in error_message:
                error_message = "Invalid or corrupted PDF file."

            failed_files.append(file.filename)
            failed_reasons.append(error_message)
    print("SUCCESS COUNT =", success_count)

    print("Saved structured data to database!")
    
    if success_count == 0:
        message = ""
    elif success_count == 1:
        message = "1 resume processed successfully"
    else:
        message = f"{success_count} resumes processed successfully"
    error = ""
        

    if failed_files:
        error = f"&error={failed_files[0]}: {failed_reasons[0]}"

    return redirect(
        f"/?message={message}{error}"
    )

@app.route("/search")
@login_required
def search():

    skill = request.args.get("skill", "")
    location = request.args.get("location", "")

    min_experience = request.args.get("experience", "")

    if (
        skill.strip() == "" and
        location.strip() == "" and
        min_experience.strip() == ""
    ):
        return """
        <html>
        <body style="font-family: Arial; margin:40px;">
            <h2>Please enter at least one search criteria.</h2>

            <br>

            <a href="/">Back</a>

        </body>
        </html>
        """

    if min_experience == "":
        min_experience = "0"

    print("MIN EXPERIENCE =", min_experience)

    conn = sqlite3.connect("candidates.db")
    cursor = conn.cursor()

    cursor.execute("""
        SELECT *
        FROM candidates
        WHERE skills LIKE ?
        AND location LIKE ?
    """, (
        f"%{skill}%",
        f"%{location}%"
    ))

    results = cursor.fetchall()
    
    filtered_results = []

    for row in results:
        try:
            candidate_experience = float(row[7])

            if candidate_experience >= float(min_experience):
                filtered_results.append(row)

        except Exception as e:
            print("ERROR =", e)
    print(f"Search completed. Results found:{len(filtered_results)}")

    app.logger.info(
        f"Search: skill='{skill}', location='{location}', experience='{min_experience}' | Results: {len(filtered_results)}"
    )

    conn.close()

    output = """
    <html>
    <head>
    <style>

    .profile-btn {
        display: inline-block;
        margin-top: 10px;
        padding: 8px 15px;
        background: #007bff;
        color: white;
        text-decoration: none;
        border-radius: 5px;
    }

    .profile-btn:hover {
        background: #0056b3;
    }

    </style>
    </head>

    <body style="font-family: Arial; margin: 40px;">
    <h1>Search Results</h1>
    """

    for row in filtered_results:

        output += f"""
        <div style="
            border:1px solid #ddd;
            padding:20px;
            margin-bottom:20px;
            border-radius:10px;
        ">

        <h2>{row[1]}</h2>

        <p><b>Skills:</b><br>{row[6]}</p>

        <p><b>Experience:</b> {row[7]} years</p>

        <p><b>Location:</b> {row[9]}</p>

        <a
            class="profile-btn"
            href="/profile/{row[0]}">
            View Profile
        </a>

        </div>
        """

    if len(filtered_results) == 0:
        output += """
        <h2>No candidates found</h2>
        """

    output += """
    </body>
    </html>
    """

    return output

@app.route("/profile/<int:id>")
@login_required
def profile(id):

    conn = sqlite3.connect("candidates.db")
    cursor = conn.cursor()

    cursor.execute(
        "SELECT * FROM candidates WHERE id = ?",
        (id,)
    )

    candidate = cursor.fetchone()

    conn.close()

    if not candidate:
        return "Candidate not found"

    return f"""
    <html>
    <body style="font-family: Arial; margin:40px;">

    <h1>{candidate[1]}</h1>

    <p>
    <b>Email:</b>
    <a href="mailto:{candidate[2]}">
        {candidate[2]}
    </a>
    </p>

    <p>
    <b>Phone:</b>
    <a href="tel:{candidate[3]}">
        {candidate[3]}
    </a>
    </p>

    <p>
    <b>LinkedIn:</b>
    <a href="https://{candidate[4]}" target="_blank">
        {candidate[4]}
    </a>
    </p>

    <p><b>File:</b> {candidate[5]}</p>

    <p><b>Skills:</b> {candidate[6]}</p>

    <p><b>Experience:</b> {candidate[7]}</p>

    <p><b>Job Title:</b> {candidate[8]}</p>

    <p><b>Location:</b> {candidate[9]}</p>

    <br>

    <a
        href="/resume/{candidate[5]}"
        style="
            display:inline-block;
            padding:8px 15px;
            background:#28a745;
            color:white;
            text-decoration:none;
            border-radius:5px;
        ">
        View Resume
    </a>

<br><br>

    <a href="/">Back</a>

    </body>
    </html>
    """

@app.route("/resume/<path:filename>")
@login_required
def view_resume(filename):

    return send_from_directory(
        app.config["UPLOAD_FOLDER"],
        filename,
        as_attachment=False
    )

@app.route("/retry-pending")
@login_required
def retry_pending():

    jobs = get_pending_retry_jobs()

    output = "<h1>Pending Retry Jobs</h1>"

    if len(jobs) == 0:

        output += "<p>No pending jobs.</p>"

    else:

        for job in jobs:

            output += f"""
            <p>
            Job ID: {job[0]}
            <br>
            File: {job[1]}
            <br>
            Status: {job[2]}
            <br><br>
            </p>
            """

    return output

@app.errorhandler(413)
def file_too_large(error):

    return redirect(
        "/?error=File is too large. Maximum allowed size is 5 MB."
    )

@app.errorhandler(429)
def too_many_requests(error):

    return redirect(
        "/?error=Too many upload requests. Please wait a minute and try again."
    )

@app.errorhandler(Exception)
def handle_exception(error):

    app.logger.exception(f"Unhandled exception: {str(error)}")

    return "An unexpected error occurred.", 500

if __name__ == "__main__":
    app.run()
